# -*- coding: utf-8 -*-
"""Логика поиска и замены текста в файлах разных типов и в именах
файлов/папок, с возможностью полного отката (backup).

Поддерживаемые цели замены:
  • Word (.docx)          — текст в параграфах, таблицах и колонтитулах
  • Excel (.xlsx)         — текст в ячейках
  • AutoCAD (.dxf)        — текстовые сущности (TEXT/MTEXT/ATTRIB) через ezdxf
  • AutoCAD (.dwg)        — нативный формат через LibreDWG (dwg2dxf → замена → dxf2dwg)
                            Библиотека LibreDWG ищется в папке «libredwg» рядом с программой.
  • PDF (.pdf)            — текст в PDF; перебор библиотек PyMuPDF → pikepdf → pypdf
                            (если одна не справилась, подключается следующая)
  • имена файлов          — переименование файлов
  • названия папок        — переименование папок

Каждая операция замены создаёт «транзакцию» с резервными копиями, что
позволяет полностью откатить последнюю операцию («вернуть прежнее название»).
"""

import os
import re
import shutil
import subprocess
import tempfile
import datetime
from pathlib import Path
from typing import List, Dict, Any, Callable, Optional


# Расширения файлов по типам
WORD_EXTS  = {".docx"}
EXCEL_EXTS = {".xlsx"}
DXF_EXTS   = {".dxf"}
DWG_EXTS   = {".dwg"}
PDF_EXTS   = {".pdf"}

# Группы расширений для фильтра переименования ИМЁН ФАЙЛОВ.
# Здесь намеренно перечислены и «старые» форматы (.doc/.xls) — при переименовании
# фильтруется само имя файла, а не его содержимое, поэтому ограничение
# на .docx/.xlsx тут не нужно.
FILENAME_EXT_GROUPS = {
    "word":  {".doc", ".docx", ".docm", ".rtf"},
    "excel": {".xls", ".xlsx", ".xlsm", ".xlsb", ".csv"},
    "dwg":   {".dwg", ".dxf"},
    "pdf":   {".pdf"},
}


def _available_pdf_libs() -> list:
    """Возвращает список доступных PDF-библиотек в порядке предпочтения.
    Замена пробуется по очереди: если одна не справилась — берётся следующая."""
    libs = []
    try:
        import fitz  # PyMuPDF  # noqa
        libs.append("pymupdf")
    except Exception:
        pass
    try:
        import pikepdf  # noqa
        libs.append("pikepdf")
    except Exception:
        pass
    try:
        import pypdf  # noqa
        libs.append("pypdf")
    except Exception:
        pass
    return libs


# ---------------------------------------------------------------------------
# Поиск LibreDWG: папка libredwg/ рядом с программой (exe или main.py)
# ---------------------------------------------------------------------------
def _libredwg_dir() -> Optional[Path]:
    """Возвращает путь к папке libredwg, если она существует рядом с программой."""
    import sys
    if getattr(sys, "frozen", False):
        base = Path(sys.executable).parent
    else:
        base = Path(__file__).parent
    candidate = base / "libredwg"
    return candidate if candidate.is_dir() else None


def _libredwg_exe(name: str) -> Optional[Path]:
    """Путь к утилите LibreDWG (dwg2dxf.exe / dxf2dwg.exe)."""
    d = _libredwg_dir()
    if d is None:
        return None
    exe = d / f"{name}.exe"
    return exe if exe.exists() else None


# ---------------------------------------------------------------------------
# Вспомогательные классы
# ---------------------------------------------------------------------------
class ReplaceResult:
    """Результат замены в одном файле/папке — узел журнала."""
    def __init__(self, path: str, kind: str):
        self.path     = path    # путь к файлу/папке
        self.kind     = kind    # 'word'|'excel'|'autocad'|'filename'|'foldername'
        self.changes  = []      # список строк-описаний замен (для дерева журнала)
        self.count    = 0       # общее число замен
        self.new_path = None    # новый путь (для переименований)
        self.error    = None    # текст ошибки, если была

    def add_change(self, description: str):
        self.changes.append(description)
        self.count += 1


class ReplaceTransaction:
    """Одна операция замены: набор резервных копий для отката."""
    def __init__(self, backup_dir: Path):
        self.backup_dir = backup_dir
        self.timestamp  = datetime.datetime.now()
        # Список действий для отката:
        #  ('content', original_path, backup_path)
        #  ('rename',  new_path, original_path)
        self.actions: list = []
        self.results: List[ReplaceResult] = []
        # Список пар (find, replace) данной операции — для журнала
        self.pairs: List[tuple] = []

    def record_content_backup(self, original_path: str, backup_path: str):
        self.actions.append(("content", original_path, backup_path))

    def record_rename(self, new_path: str, original_path: str):
        self.actions.append(("rename", new_path, original_path))


def _backup_file_content(path: Path, backup_dir: Path, counter: list) -> Path:
    """Копирует файл в backup_dir с уникальным именем. Возвращает путь копии."""
    backup_dir.mkdir(parents=True, exist_ok=True)
    counter[0] += 1
    bpath = backup_dir / f"{counter[0]:05d}_{path.name}"
    shutil.copy2(path, bpath)
    return bpath


# ---------------------------------------------------------------------------
# Замена в Word (.docx)
# ---------------------------------------------------------------------------
def replace_in_docx(path: Path, find: str, replace: str) -> ReplaceResult:
    res = ReplaceResult(str(path), "word")
    try:
        import docx
        doc = docx.Document(str(path))

        def process_paragraph(p):
            if find in p.text:
                full  = p.text
                n     = full.count(find)
                new_t = full.replace(find, replace)
                for i, run in enumerate(p.runs):
                    run.text = new_t if i == 0 else ""
                return n
            return 0

        def process_tables(tables, where):
            """Обходит таблицы (и вложенные в ячейки) указанной части документа."""
            total = 0
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            total += process_paragraph(p)
                        # таблицы могут быть вложены в ячейку
                        total += process_tables(cell.tables, where)
            return total

        # --- Тело документа ---
        body_p = sum(process_paragraph(p) for p in doc.paragraphs)
        if body_p:
            res.add_change(f"Параграф: «{find}» → «{replace}» ({body_p} шт.)")
        body_t = process_tables(doc.tables, "body")
        if body_t:
            res.add_change(f"Таблица: «{find}» → «{replace}» ({body_t} шт.)")

        # --- Колонтитулы каждого раздела (верхние и нижние, включая
        #     отдельные колонтитулы первой и чётных страниц) ---
        hf_parts = (
            ("Верхний колонтитул",                "header"),
            ("Нижний колонтитул",                 "footer"),
            ("Верхний колонтитул (первая стр.)",  "first_page_header"),
            ("Нижний колонтитул (первая стр.)",   "first_page_footer"),
            ("Верхний колонтитул (чётные стр.)",  "even_page_header"),
            ("Нижний колонтитул (чётные стр.)",   "even_page_footer"),
        )
        for section in doc.sections:
            for label, attr in hf_parts:
                part = getattr(section, attr, None)
                if part is None:
                    continue
                # Пропускаем «связанные» колонтитулы (наследуют предыдущий
                # раздел) — иначе посчитаем замену повторно.
                if getattr(part, "is_linked_to_previous", False):
                    continue
                cnt = sum(process_paragraph(p) for p in part.paragraphs)
                cnt += process_tables(part.tables, attr)
                if cnt:
                    res.add_change(
                        f"{label}: «{find}» → «{replace}» ({cnt} шт.)")

        if res.count > 0:
            doc.save(str(path))
    except Exception as e:
        res.error = str(e)
    return res


# ---------------------------------------------------------------------------
# Замена в Excel (.xlsx)
# ---------------------------------------------------------------------------
def replace_in_xlsx(path: Path, find: str, replace: str) -> ReplaceResult:
    res = ReplaceResult(str(path), "excel")
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path))
        for ws in wb.worksheets:
            sheet_title = ws.title  # запоминаем до возможного переименования

            # ── Ячейки ────────────────────────────────────────────────────
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str) and find in cell.value:
                        n          = cell.value.count(find)
                        cell.value = cell.value.replace(find, replace)
                        res.add_change(
                            f"Лист «{sheet_title}», {cell.coordinate}: "
                            f"«{find}» → «{replace}» ({n} шт.)")

            # ── Название листа ────────────────────────────────────────────
            if find in ws.title:
                new_title = ws.title.replace(find, replace)
                ws.title  = new_title
                res.add_change(
                    f"Название листа: «{sheet_title}» → «{new_title}»")

            # ── Колонтитулы (odd / even / first) ─────────────────────────
            for hf_attr in ("oddHeader", "oddFooter",
                            "evenHeader", "evenFooter",
                            "firstHeader", "firstFooter"):
                hf = getattr(ws.HeaderFooter, hf_attr, None)
                if hf is None:
                    continue
                changed = False
                for part_attr in ("left", "center", "right"):
                    part = getattr(hf, part_attr, None)
                    if part and part.text and find in part.text:
                        n          = part.text.count(find)
                        part.text  = part.text.replace(find, replace)
                        changed    = True
                        label = {"oddHeader": "Верхний колонтитул",
                                 "oddFooter": "Нижний колонтитул",
                                 "evenHeader": "Верхний колонтитул (чётные)",
                                 "evenFooter": "Нижний колонтитул (чётные)",
                                 "firstHeader": "Верхний колонтитул (первая)",
                                 "firstFooter": "Нижний колонтитул (первая)",
                                 }.get(hf_attr, hf_attr)
                        res.add_change(
                            f"Лист «{sheet_title}», {label} ({part_attr}): "
                            f"«{find}» → «{replace}» ({n} шт.)")

        if res.count > 0:
            wb.save(str(path))
    except Exception as e:
        res.error = str(e)
    return res


# ---------------------------------------------------------------------------
# Замена в AutoCAD: ядро — работа с ezdxf-документом (DXF или временный DXF)
# ---------------------------------------------------------------------------
def _replace_in_dxf_doc(doc, find: str, replace: str,
                        res: "ReplaceResult") -> int:
    """Заменяет текст во всех текстовых сущностях открытого ezdxf-документа.
    Обрабатывает TEXT, MTEXT, ATTRIB (в т.ч. атрибуты INSERT) во всех
    layouts и блоках. Возвращает число замен."""
    total = 0

    def process_space(space):
        nonlocal total
        for entity in space:
            try:
                # Защита: иногда при обходе блоков попадаются не-сущности
                # (например, имена блоков как строки) — пропускаем их.
                if not hasattr(entity, "dxftype"):
                    continue
                dxftype = entity.dxftype()
                if dxftype == "TEXT":
                    val = entity.dxf.text or ""
                    if find in val:
                        n = val.count(find)
                        entity.dxf.text = val.replace(find, replace)
                        total += n
                        res.add_change(f"TEXT: «{find}» → «{replace}» ({n} шт.)")
                elif dxftype == "MTEXT":
                    val = entity.text or ""
                    if find in val:
                        n = val.count(find)
                        entity.text = val.replace(find, replace)
                        total += n
                        res.add_change(f"MTEXT: «{find}» → «{replace}» ({n} шт.)")
                elif dxftype == "ATTRIB":
                    val = entity.dxf.text or ""
                    if find in val:
                        n = val.count(find)
                        entity.dxf.text = val.replace(find, replace)
                        total += n
                        res.add_change(f"ATTRIB: «{find}» → «{replace}» ({n} шт.)")
                elif dxftype == "INSERT":
                    for attrib in entity.attribs:
                        val = attrib.dxf.text or ""
                        if find in val:
                            n = val.count(find)
                            attrib.dxf.text = val.replace(find, replace)
                            total += n
                            res.add_change(
                                f"Атрибут блока: «{find}» → «{replace}» ({n} шт.)")
            except Exception:
                continue

    for layout in doc.layouts:
        try:
            process_space(layout)
        except Exception:
            continue
    for block in doc.blocks:
        try:
            process_space(block)
        except Exception:
            continue
    return total


def replace_in_dxf(path: Path, find: str, replace: str) -> ReplaceResult:
    """Замена в DXF. Основной путь — ezdxf; если он не справился (ошибка
    чтения/записи), пробуем нативную библиотеку libdxfrw (C/C++) через её
    DLL-обёртку, если она собрана и доступна."""
    res = ReplaceResult(str(path), "autocad")
    ezdxf_error = None
    try:
        import ezdxf
        doc   = ezdxf.readfile(str(path))
        total = _replace_in_dxf_doc(doc, find, replace, res)
        if total > 0:
            doc.saveas(str(path))
        return res
    except Exception as e:
        ezdxf_error = str(e)
        # Сбросим частичный результат перед попыткой запасной библиотеки
        res = ReplaceResult(str(path), "autocad")

    # --- Запасной путь №1: нативная libdxfrw ---
    libdxfrw_error = None
    try:
        import dxfrw_native
        if dxfrw_native.is_available():
            tmp_out = path.with_suffix(".dxf.dxfrw_out")
            n = dxfrw_native.replace_text_in_dxf(str(path), str(tmp_out),
                                                 find, replace)
            if n > 0 and tmp_out.exists():
                shutil.move(str(tmp_out), str(path))
                res.add_change(
                    f"DXF (libdxfrw): «{find}» → «{replace}» ({n} шт.)")
                return res
            if tmp_out.exists():
                tmp_out.unlink()
            # 0 замен — текста нет; считаем обработанным
            return res
    except Exception as e2:
        libdxfrw_error = str(e2)

    # --- Запасной путь №2: ACadSharp (.NET-утилита) ---
    try:
        import acadsharp_native
        if acadsharp_native.is_available():
            tmp_out = path.with_suffix(".dxf.acad_out")
            n = acadsharp_native.replace_text(str(path), str(tmp_out),
                                              find, replace)
            if n > 0 and tmp_out.exists():
                shutil.move(str(tmp_out), str(path))
                res.add_change(
                    f"DXF (ACadSharp): «{find}» → «{replace}» ({n} шт.)")
                return res
            if tmp_out.exists():
                tmp_out.unlink()
            return res
    except Exception as e3:
        res.error = (f"DXF: ezdxf — {ezdxf_error}; "
                     f"libdxfrw — {libdxfrw_error or 'недоступна'}; "
                     f"ACadSharp — {e3}")
        return res

    # Все запасные пути недоступны — сообщаем исходную ошибку ezdxf
    suffix = ""
    if libdxfrw_error:
        suffix = f"; libdxfrw — {libdxfrw_error}"
    res.error = f"DXF: {ezdxf_error}{suffix}"
    return res


# ---------------------------------------------------------------------------
# Замена в PDF (несколько библиотек с автоматическим фолбэком)
# ---------------------------------------------------------------------------
def _pdf_replace_pymupdf(path: Path, find: str, replace: str, res) -> int:
    """Замена через PyMuPDF (fitz): находит вхождения текста, закрашивает их
    и впечатывает новый текст на то же место. Лучший визуальный результат."""
    import fitz
    doc = fitz.open(str(path))
    total = 0
    try:
        for page in doc:
            rects = page.search_for(find)
            if not rects:
                continue
            for r in rects:
                # Закрашиваем старый текст фоном страницы (обычно белый)
                page.add_redact_annot(r, fill=(1, 1, 1))
            page.apply_redactions()
            for r in rects:
                # Впечатываем новый текст в левый нижний угол прежнего прямоугольника
                fs = max(6, min(12, r.height * 0.8))
                page.insert_text((r.x0, r.y1 - r.height * 0.18),
                                 replace, fontsize=fs, color=(0, 0, 0))
                total += 1
        if total:
            # incremental=False — полностью пересохраняем файл
            tmp = path.with_suffix(".pdf.tmp")
            doc.save(str(tmp), garbage=4, deflate=True)
            doc.close()
            shutil.move(str(tmp), str(path))
            res.add_change(f"PDF (PyMuPDF): «{find}» → «{replace}» ({total} шт.)")
            return total
    finally:
        try:
            doc.close()
        except Exception:
            pass
    return 0


def _pdf_replace_stream(path: Path, find: str, replace: str, res, lib: str) -> int:
    """Замена на уровне потоков содержимого страницы (pikepdf или pypdf).
    Работает для PDF, где текст хранится как простые строки в content stream
    (типичный случай для файлов, сгенерированных программами/Excel/Word).
    Не перерисовывает раскладку — просто подменяет байты строк, что подходит,
    когда длина текста меняется незначительно (например, шифры)."""
    fb = find.encode("latin-1", "ignore")
    rb = replace.encode("latin-1", "ignore")
    if not fb:
        return 0
    total = 0

    if lib == "pikepdf":
        import pikepdf
        pdf = pikepdf.open(str(path), allow_overwriting_input=True)
        try:
            for page in pdf.pages:
                if "/Contents" not in page:
                    continue
                contents = page.Contents
                streams = contents if isinstance(contents, pikepdf.Array) else [contents]
                for stream in streams:
                    try:
                        data = stream.read_bytes()
                    except Exception:
                        continue
                    if fb in data:
                        cnt = data.count(fb)
                        stream.write(data.replace(fb, rb))
                        total += cnt
            if total:
                pdf.save(str(path))
                res.add_change(f"PDF (pikepdf): «{find}» → «{replace}» ({total} шт.)")
        finally:
            pdf.close()
        return total

    if lib == "pypdf":
        import pypdf
        from pypdf.generic import DecodedStreamObject, EncodedStreamObject
        reader = pypdf.PdfReader(str(path))
        writer = pypdf.PdfWriter()
        for page in reader.pages:
            try:
                contents = page.get_contents()
                if contents is not None:
                    data = contents.get_data()
                    if fb in data:
                        cnt = data.count(fb)
                        contents.set_data(data.replace(fb, rb))
                        page.replace_contents(contents)
                        total += cnt
            except Exception:
                pass
            writer.add_page(page)
        if total:
            with open(path, "wb") as f:
                writer.write(f)
            res.add_change(f"PDF (pypdf): «{find}» → «{replace}» ({total} шт.)")
        return total

    return 0


def replace_in_pdf(path: Path, find: str, replace: str) -> ReplaceResult:
    """Замена текста в PDF с автоматическим перебором библиотек.

    Порядок попыток (если одна не справилась — подключается следующая):
      1) PyMuPDF (fitz)  — лучший результат, перерисовывает текст;
      2) pikepdf         — правка потоков содержимого;
      3) pypdf           — правка потоков содержимого.

    Возвращает результат первой библиотеки, выполнившей хотя бы одну замену.
    Если ни одна не справилась, но текст в файле есть — сообщает об этом.
    """
    res = ReplaceResult(str(path), "pdf")
    libs = _available_pdf_libs()
    if not libs:
        res.error = ("PDF: не установлена ни одна библиотека "
                     "(pymupdf, pikepdf или pypdf).")
        return res

    errors = []
    for lib in libs:
        # Каждую попытку делаем на свежей копии — если библиотека частично
        # испортила файл, следующая стартует с оригинала.
        backup = None
        try:
            backup = path.with_suffix(".pdf.bak_try")
            shutil.copy2(path, backup)
            if lib == "pymupdf":
                n = _pdf_replace_pymupdf(path, find, replace, res)
            else:
                n = _pdf_replace_stream(path, find, replace, res, lib)
            if n > 0:
                if backup and backup.exists():
                    backup.unlink()
                return res
            # 0 замен — восстановим оригинал и пробуем следующую библиотеку
            if backup and backup.exists():
                shutil.move(str(backup), str(path))
                backup = None
        except Exception as e:
            errors.append(f"{lib}: {e}")
            # откатываем возможную порчу и пробуем следующую
            if backup and backup.exists():
                shutil.move(str(backup), str(path))
                backup = None
            continue
        finally:
            if backup and backup.exists():
                try:
                    backup.unlink()
                except Exception:
                    pass

    if errors:
        res.error = "PDF: ни одна библиотека не справилась. " + "; ".join(errors)
    # если ошибок нет и замен нет — значит, искомого текста в файле просто нет
    return res


# ---------------------------------------------------------------------------
# Замена в DWG через LibreDWG (dwg2dxf → замена → dxf2dwg)
# ---------------------------------------------------------------------------
def _run_libredwg(exe_name: str, *args) -> subprocess.CompletedProcess:
    """Запускает утилиту LibreDWG с заданными аргументами.
    Добавляет папку libredwg/ в PATH, чтобы подхватились зависимые DLL."""
    exe = _libredwg_exe(exe_name)
    if exe is None:
        raise FileNotFoundError(
            f"Утилита {exe_name}.exe не найдена. "
            f"Убедитесь, что папка libredwg\\ находится рядом с программой.")

    env = os.environ.copy()
    lib_dir = str(exe.parent)
    env["PATH"] = lib_dir + os.pathsep + env.get("PATH", "")

    # На Windows прячем чёрное консольное окно (приложение оконное).
    creationflags = 0
    if os.name == "nt":
        creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)

    result = subprocess.run(
        [str(exe)] + list(args),
        capture_output=True,
        env=env,
        timeout=60,
        creationflags=creationflags,
    )
    return result


def replace_in_dwg(path: Path, find: str, replace: str) -> ReplaceResult:
    """Замена в DWG. Пути (по очереди, «не сработал один — пробуем другой»):

      A. ACadSharp (.NET-утилита, если собрана) — читает и пишет DWG напрямую,
         без промежуточной конвертации. Предпочтительно, когда доступна.
      B. LibreDWG round-trip: dwg2dxf → замена (ezdxf, при сбое libdxfrw) → dxf2dwg.

    Требует для пути B папку libredwg\\ рядом с программой.
    """
    res = ReplaceResult(str(path), "autocad")

    # ── Путь A: ACadSharp напрямую (DWG→DWG) ─────────────────────────────
    try:
        import acadsharp_native
        if acadsharp_native.is_available():
            tmp_out = path.with_suffix(".dwg.acad_out")
            try:
                n = acadsharp_native.replace_text(str(path), str(tmp_out),
                                                  find, replace)
                if n > 0 and tmp_out.exists():
                    shutil.move(str(tmp_out), str(path))
                    res.add_change(
                        f"DWG (ACadSharp): «{find}» → «{replace}» ({n} шт.)")
                    return res
                if n == 0:
                    if tmp_out.exists():
                        tmp_out.unlink()
                    return res  # текста нет — обработано
                if tmp_out.exists():
                    tmp_out.unlink()
            except Exception:
                # ACadSharp не справился — переходим к пути B (LibreDWG)
                if tmp_out.exists():
                    try:
                        tmp_out.unlink()
                    except Exception:
                        pass
    except Exception:
        pass  # модуль недоступен — путь B

    # ── Путь B: LibreDWG round-trip ──────────────────────────────────────
    # Проверяем наличие ezdxf
    try:
        import ezdxf  # noqa
    except ImportError:
        res.error = "DWG: модуль ezdxf не установлен (pip install ezdxf)"
        return res

    # Проверяем наличие утилит LibreDWG
    if _libredwg_exe("dwg2dxf") is None or _libredwg_exe("dxf2dwg") is None:
        res.error = (
            "DWG: утилиты LibreDWG не найдены. "
            "Поместите папку libredwg\\ рядом с программой."
        )
        return res

    tmp_dir = Path(tempfile.mkdtemp(prefix="nept_dwg_"))
    try:
        # ── Шаг 1: DWG → DXF ──────────────────────────────────────────────
        # ВАЖНО: у dwg2dxf флаг -o задаёт ИМЯ ВЫХОДНОГО ФАЙЛА (не директорию).
        tmp_dxf = tmp_dir / (path.stem + ".dxf")
        r1 = _run_libredwg("dwg2dxf", "-o", str(tmp_dxf), str(path))
        if not tmp_dxf.exists():
            stderr = r1.stderr.decode("utf-8", errors="replace").strip()
            res.error = f"DWG→DXF: конвертация не удалась. {stderr}"
            return res

        # ── Шаг 2: замена текста в DXF ─────────────────────────────────────
        tmp_dxf_out = tmp_dir / (path.stem + "_out.dxf")
        total = 0
        ezdxf_ok = False
        try:
            import ezdxf
            try:
                doc = ezdxf.readfile(str(tmp_dxf))
            except Exception:
                # DXF от LibreDWG бывает «не совсем чистым» (ошибки вида
                # «256 is not a valid ColumnType») — пробуем режим восстановления.
                from ezdxf import recover
                doc, auditor = recover.readfile(str(tmp_dxf))
            total = _replace_in_dxf_doc(doc, find, replace, res)
            if total > 0:
                doc.saveas(str(tmp_dxf_out))
            ezdxf_ok = True
        except Exception as ez_err:
            # Запасной путь: нативная libdxfrw обрабатывает промежуточный DXF
            try:
                import dxfrw_native
                if dxfrw_native.is_available():
                    n = dxfrw_native.replace_text_in_dxf(
                        str(tmp_dxf), str(tmp_dxf_out), find, replace)
                    total = n
                    if n > 0:
                        res.add_change(
                            f"DWG (libdxfrw): «{find}» → «{replace}» ({n} шт.)")
                else:
                    res.error = f"DWG: ezdxf — {ez_err}; libdxfrw не собрана"
                    return res
            except Exception as nat_err:
                res.error = f"DWG: ezdxf — {ez_err}; libdxfrw — {nat_err}"
                return res

        if total == 0:
            return res  # нечего сохранять
        if not tmp_dxf_out.exists():
            return res

        # ── Шаг 3: DXF → DWG ──────────────────────────────────────────────
        # У dxf2dwg флаг -o также задаёт имя выходного файла.
        tmp_dwg_out = tmp_dir / (path.stem + "_out.dwg")
        r2 = _run_libredwg("dxf2dwg", "-y", "-o", str(tmp_dwg_out), str(tmp_dxf_out))
        if not tmp_dwg_out.exists():
            stderr = r2.stderr.decode("utf-8", errors="replace").strip()
            res.error = f"DXF→DWG: конвертация не удалась. {stderr}"
            return res

        # Перезаписываем оригинальный DWG результатом
        shutil.copy2(tmp_dwg_out, path)

    except Exception as e:
        res.error = f"DWG: {e}"
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return res


# ---------------------------------------------------------------------------
# Главная функция замены
# ---------------------------------------------------------------------------
def run_replace(
    root_dir: str,
    find: str,
    replace: str,
    targets: Dict[str, bool],
    backup_root: Path,
    progress: Optional[Callable[[str], None]] = None,
    filename_exts: Optional[set] = None,
    _tx: Optional["ReplaceTransaction"] = None,
    _counter: Optional[list] = None,
) -> ReplaceTransaction:
    """Выполняет замену во всех файлах папки (рекурсивно) согласно целям.

    targets: {"word": bool, "excel": bool, "autocad": bool,
              "filename": bool, "foldername": bool}
    filename_exts: множество расширений (в нижнем регистре, с точкой), которыми
              ограничивается переименование ИМЁН ФАЙЛОВ. Если None или пусто —
              переименовываются файлы с любым расширением (прежнее поведение).
    _tx, _counter: служебные параметры для накопления нескольких пар
              «найти/заменить» в одну транзакцию (см. run_replace_multi).
    Возвращает транзакцию с результатами и данными для отката.
    """
    root = Path(root_dir)
    ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = backup_root / f"replace_{ts}"
    if _tx is not None:
        tx = _tx
        backup_dir = tx.backup_dir
    else:
        tx = ReplaceTransaction(backup_dir)
    counter = _counter if _counter is not None else [0]

    if filename_exts:
        filename_exts = {e.lower() for e in filename_exts}

    def log(msg):
        if progress:
            try:
                progress(msg)
            except Exception:
                pass

    if not find:
        return tx

    # ── 1) Замена СОДЕРЖИМОГО файлов ─────────────────────────────────────
    content_targets = []
    if targets.get("word"):
        content_targets.append(("word",    WORD_EXTS))
    if targets.get("excel"):
        content_targets.append(("excel",   EXCEL_EXTS))
    if targets.get("autocad"):
        content_targets.append(("autocad", DXF_EXTS | DWG_EXTS))
    if targets.get("pdf"):
        content_targets.append(("pdf", PDF_EXTS))

    if content_targets:
        for dirpath, dirnames, filenames in os.walk(root):
            for fname in filenames:
                ext   = Path(fname).suffix.lower()
                fpath = Path(dirpath) / fname
                for kind, exts in content_targets:
                    if ext not in exts:
                        continue
                    log(f"Обработка: {fpath}")
                    try:
                        bpath = _backup_file_content(fpath, backup_dir, counter)
                    except Exception:
                        bpath = None

                    if kind == "word":
                        r = replace_in_docx(fpath, find, replace)
                    elif kind == "excel":
                        r = replace_in_xlsx(fpath, find, replace)
                    elif kind == "autocad":
                        if ext in DWG_EXTS:
                            r = replace_in_dwg(fpath, find, replace)
                        else:
                            r = replace_in_dxf(fpath, find, replace)
                    elif kind == "pdf":
                        r = replace_in_pdf(fpath, find, replace)

                    if r.count > 0 and bpath is not None:
                        tx.record_content_backup(str(fpath), str(bpath))
                    if r.count > 0 or r.error:
                        tx.results.append(r)

    # ── 2) Переименование ФАЙЛОВ ──────────────────────────────────────────
    if targets.get("filename"):
        for dirpath, dirnames, filenames in os.walk(root):
            for fname in filenames:
                # Фильтр по расширению (если задан список)
                if filename_exts:
                    if Path(fname).suffix.lower() not in filename_exts:
                        continue
                if find in fname:
                    new_name = fname.replace(find, replace)
                    old_path = Path(dirpath) / fname
                    new_path = Path(dirpath) / new_name
                    if new_path.exists():
                        r = ReplaceResult(str(old_path), "filename")
                        r.error = f"Файл «{new_name}» уже существует — пропущено"
                        tx.results.append(r)
                        continue
                    try:
                        old_path.rename(new_path)
                        r          = ReplaceResult(str(old_path), "filename")
                        r.new_path = str(new_path)
                        r.add_change(f"Имя файла: «{fname}» → «{new_name}»")
                        tx.record_rename(str(new_path), str(old_path))
                        tx.results.append(r)
                        log(f"Переименован файл: {fname} → {new_name}")
                    except Exception as e:
                        r = ReplaceResult(str(old_path), "filename")
                        r.error = str(e)
                        tx.results.append(r)

    # ── 3) Переименование ПАПОК (снизу вверх) ─────────────────────────────
    if targets.get("foldername"):
        all_dirs = []
        for dirpath, dirnames, filenames in os.walk(root):
            all_dirs.append(dirpath)
        all_dirs.sort(key=lambda p: p.count(os.sep), reverse=True)
        for dirpath in all_dirs:
            d = Path(dirpath)
            if d == root:
                continue
            if find in d.name:
                new_name = d.name.replace(find, replace)
                new_path = d.parent / new_name
                if new_path.exists():
                    r = ReplaceResult(str(d), "foldername")
                    r.error = f"Папка «{new_name}» уже существует — пропущено"
                    tx.results.append(r)
                    continue
                try:
                    d.rename(new_path)
                    r          = ReplaceResult(str(d), "foldername")
                    r.new_path = str(new_path)
                    r.add_change(f"Имя папки: «{d.name}» → «{new_name}»")
                    tx.record_rename(str(new_path), str(d))
                    tx.results.append(r)
                    log(f"Переименована папка: {d.name} → {new_name}")
                except Exception as e:
                    r = ReplaceResult(str(d), "foldername")
                    r.error = str(e)
                    tx.results.append(r)

    return tx


# ---------------------------------------------------------------------------
# Замена по НЕСКОЛЬКИМ парам «найти/заменить» в одну транзакцию
# ---------------------------------------------------------------------------
def run_replace_multi(
    root_dir: str,
    pairs: List[tuple],
    targets: Dict[str, bool],
    backup_root: Path,
    progress: Optional[Callable[[str], None]] = None,
    filename_exts: Optional[set] = None,
) -> ReplaceTransaction:
    """Выполняет несколько замен подряд, накапливая результаты и резервные
    копии в ОДНУ транзакцию (чтобы откат отменял всю операцию целиком).

    pairs: список кортежей (find, replace). Пустые find пропускаются.
    Остальные параметры — как у run_replace.
    """
    # Единая транзакция и общий счётчик резервных копий на все пары
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = backup_root / f"replace_{ts}"
    tx = ReplaceTransaction(backup_dir)
    counter = [0]
    # Запоминаем все пары для отображения в журнале
    tx.pairs = [(f, r) for f, r in pairs if f]

    for find, replace in pairs:
        if not find:
            continue
        run_replace(
            root_dir, find, replace, targets, backup_root,
            progress=progress, filename_exts=filename_exts,
            _tx=tx, _counter=counter,
        )
    return tx


# ---------------------------------------------------------------------------
# Замена по списку ЗАДАЧ — у каждой свои текст, цели и расширения
# ---------------------------------------------------------------------------
def run_replace_tasks(
    root_dir: str,
    tasks: List[dict],
    backup_root: Path,
    progress: Optional[Callable[[str], None]] = None,
) -> ReplaceTransaction:
    """Выполняет несколько ЗАДАЧ замены, у каждой собственная конфигурация.

    tasks: список словарей вида
        {
          "find": str, "replace": str,
          "targets": {"word":bool, "excel":bool, "autocad":bool,
                      "pdf":bool, "filename":bool, "foldername":bool},
          "filename_exts": set[str] | None,   # для переименования имён файлов
        }
    Все задачи накапливаются в ОДНУ транзакцию ради единого отката.
    """
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_dir = backup_root / f"replace_{ts}"
    tx = ReplaceTransaction(backup_dir)
    counter = [0]
    tx.pairs = [(t.get("find", ""), t.get("replace", ""))
                for t in tasks if t.get("find")]

    for t in tasks:
        find = t.get("find", "")
        if not find:
            continue
        replace = t.get("replace", "")
        targets = t.get("targets", {})
        filename_exts = t.get("filename_exts") or None
        run_replace(
            root_dir, find, replace, targets, backup_root,
            progress=progress, filename_exts=filename_exts,
            _tx=tx, _counter=counter,
        )
    return tx
def rollback(tx: ReplaceTransaction,
             progress: Optional[Callable[[str], None]] = None) -> Dict[str, int]:
    """Откатывает транзакцию: восстанавливает содержимое из backup и
    возвращает прежние имена файлов/папок.

    Возвращает {"content": N, "renames": M, "errors": K}.
    """
    stats = {"content": 0, "renames": 0, "errors": 0}

    def log(msg):
        if progress:
            try:
                progress(msg)
            except Exception:
                pass

    for action in reversed(tx.actions):
        try:
            if action[0] == "content":
                _, original_path, backup_path = action
                if Path(backup_path).exists():
                    shutil.copy2(backup_path, original_path)
                    stats["content"] += 1
                    log(f"Восстановлено: {original_path}")
            elif action[0] == "rename":
                _, new_path, original_path = action
                if Path(new_path).exists():
                    Path(new_path).rename(original_path)
                    stats["renames"] += 1
                    log(f"Возвращено имя: {original_path}")
        except Exception:
            stats["errors"] += 1
    return stats
